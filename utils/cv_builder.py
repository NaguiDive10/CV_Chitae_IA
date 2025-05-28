# utils/cv_builder.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_docx_from_text(cv_text: str, filename: str = "cv_adapte.docx") -> str:
    doc = Document()
    
    # Styles par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Parcours ligne par ligne
    for line in cv_text.strip().split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph("")  # saut de ligne
        elif line.isupper():
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.runs[0].bold = True
        elif line.endswith(":"):
            para = doc.add_paragraph(line)
            para.runs[0].bold = True
        else:
            doc.add_paragraph(line)

    output_path = os.path.join("uploads", filename)
    os.makedirs("uploads", exist_ok=True)
    doc.save(output_path)
    return output_path
